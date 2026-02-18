from docx import Document


def main() -> None:
    doc = Document()

    title = "CarGo / Carl Rent autókölcsönző rendszer fejlesztői dokumentáció"
    doc.add_heading(title, 0)

    # Tartalomjegyzék
    doc.add_heading("Tartalomjegyzék", level=1)
    contents = [
        "1. Tartalomjegyzék",
        "2. Célok",
        "   2.1 Felhasználói célok",
        "   2.2 Kényelmes keresési és foglalási funkciók",
        "   2.3 Biztonságos és megbízható felhasználói fiókok",
        "   2.4 Kommunikáció és értesítések",
        "   2.5 Üzleti célok",
        "   2.6 Növekvő ügyfélkör és piacra lépés",
        "   2.7 Gazdasági hatékonyság és költségcsökkentés",
        "   2.8 Adatbiztonság és megfelelőség",
        "3. Követelmények",
        "   3.1 Funkcionális követelmények",
        "   3.2 Foglalási rendszer",
        "   3.3 Felhasználói fiókok kezelése",
        "   3.4 Adminisztrációs felület",
        "   3.5 Kommunikáció és értesítések",
        "   3.6 Biztonság",
        "   3.7 Megbízhatóság és rendelkezésre állás",
        "   3.8 Használhatóság",
        "4. Fejlesztői környezet",
        "   4.1 Backend és webes technológiák",
        "   4.2 Admin kliens (WPF / .NET)",
        "5. Kód dokumentáció",
        "6. Fejlesztési lehetőségek",
        "7. Tesztelés",
    ]
    for line in contents:
        doc.add_paragraph(line, style="List Bullet")

    # 2. Célok
    doc.add_heading("2. Célok", level=1)
    intro = (
        "A projekt célja egy webalapú autókölcsönző rendszer, amely egységesen kezeli az "
        "ügyféloldali foglalásokat, a felhasználói fiókokat, az autóflottát és az adminisztrációt. "
        "Három fő komponens: Node.js + MySQL backend, publikus webfelület (HTML/CSS/JS), "
        "valamint WPF alapú admin kliens."
    )
    doc.add_paragraph(intro)

    # 2.1 Felhasználói célok
    doc.add_heading("2.1 Felhasználói célok", level=2)
    bullets = [
        "Intuitív UI: egyértelmű navigáció az autólista, foglalás, regisztráció, "
        "bejelentkezés, saját foglalások között.",
        "Reszponzív design és gyors betöltés.",
        "Átlátható státuszok: minden műveletnél világos visszajelzés.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    # 2.2 Kényelmes keresési és foglalási funkciók
    doc.add_heading("2.2 Kényelmes keresési és foglalási funkciók", level=2)
    bullets = [
        "Szűrés és keresés: kategória, ár, üzemanyag, kulcsszó.",
        "Egyszerű foglalási folyamat: dátum kiválasztás, feltételek elfogadása, visszaigazolás.",
        "Foglalás követése: előzmények és státusz megtekintése.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    # 2.3 Biztonságos és megbízható felhasználói fiókok
    doc.add_heading("2.3 Biztonságos és megbízható felhasználói fiókok", level=2)
    bullets = [
        "Jelszavak hash-elve (bcryptjs), email verifikáció és jelszó-visszaállítás.",
        "Adatvédelem: személyes adatok minimalizálása és védelme.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    # 2.4 Kommunikáció és értesítések
    doc.add_heading("2.4 Kommunikáció és értesítések", level=2)
    bullets = [
        "Automatikus email értesítések regisztrációról, foglalásról, módosításról.",
        "Kapcsolati űrlap: üzenetek mentése, admin státuszkezelés (új/olvasott/válaszolt).",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    # 2.5 Üzleti célok
    doc.add_heading("2.5 Üzleti célok", level=2)
    bullets = [
        "Flotta menedzsment: autók létrehozása/módosítása/törlése, árak és elérhetőség kezelése.",
        "Foglalás-kezelés: admin jóváhagyás, státuszváltás, törlés.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    # 2.6 Növekvő ügyfélkör és piacra lépés
    doc.add_heading("2.6 Növekvő ügyfélkör és piacra lépés", level=2)
    doc.add_paragraph(
        "Professzionális online jelenlét, skálázható architektúra.", style="List Bullet"
    )

    # 2.7 Gazdasági hatékonyság és költségcsökkentés
    doc.add_heading("2.7 Gazdasági hatékonyság és költségcsökkentés", level=2)
    bullets = [
        "Automatizált folyamatok, kevesebb manuális adminisztráció.",
        "Adatvezérelt döntések: statisztikák, riportok.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    # 2.8 Adatbiztonság és megfelelőség
    doc.add_heading("2.8 Adatbiztonság és megfelelőség", level=2)
    bullets = [
        "GDPR alapelvek figyelembe vétele, HTTPS éles környezetben.",
        "Session alapú admin védelem, opcionális CSRF védelem.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    # 3. Követelmények
    doc.add_heading("3. Követelmények", level=1)

    doc.add_heading("3.1 Funkcionális követelmények", level=2)
    bullets = [
        "Autólista: márka, modell, évjárat, ár/nap, üzemanyag, ülésszám, kép, elérhetőség.",
        "Szűrés/rendezés/keresés: ár, kategória, kulcsszó alapján.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("3.2 Foglalási rendszer", level=2)
    bullets = [
        "Időalapú foglalás, ütközés-ellenőrzés.",
        "Email visszaigazolás, státusz: pending/confirmed/completed/cancelled.",
        "Módosítás/törlés szabályok szerint, feltételek elfogadása.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("3.3 Felhasználói fiókok kezelése", level=2)
    bullets = [
        "Regisztráció, bejelentkezés, jelszó-visszaállítás.",
        "Profiladatok szerkesztése, értesítési preferenciák.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("3.4 Adminisztrációs felület", level=2)
    bullets = [
        "Admin login szükséges.",
        "Autók CRUD, foglalások státuszkezelése, felhasználók módosítása/törlése.",
        "Statisztikák és riportok (admin/stats).",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("3.5 Kommunikáció és értesítések", level=2)
    bullets = [
        "Email értesítések minden fő eseményről.",
        "Ügyfélszolgálati űrlap, üzenet-státuszok.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("3.6 Biztonság", level=2)
    bullets = [
        "Jelszóhash, paraméterezett SQL, XSS/CSRF védelem (bővíthető), HTTPS.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("3.7 Megbízhatóság és rendelkezésre állás", level=2)
    bullets = [
        "Cél 99,5% rendelkezésre állás; MySQL mentések, process manager (PM2/Docker).",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("3.8 Használhatóság", level=2)
    bullets = [
        "WCAG-alapú akadálymentesség, többnyelvűségre felkészítés.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    # 4. Fejlesztői környezet
    doc.add_heading("4. Fejlesztői környezet", level=1)

    doc.add_heading("4.1 Backend és webes technológiák", level=2)
    bullets = [
        "Node.js + Express, MySQL (utils/database.js), dotenv konfiguráció.",
        (
            "Csomagok: express, mysql2, bcryptjs, jsonwebtoken, cors, helmet, "
            "express-rate-limit, multer, nodemailer, express-session, express-validator."
        ),
        "Futtatás: npm install, .env beállítás, npm run dev vagy npm start.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("4.2 Admin kliens (WPF / .NET 6)", level=2)
    bullets = [
        "Projekt: AdminClient (MainWindow, AdminApiService, Models, Helpers/AppConfig).",
        "Konfig: appsettings.json – ApiBaseUrl a backend /api végpontjára mutat.",
        "Futtatás: Visual Studio, .NET 6, WPF.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    # 5. Kód dokumentáció (fő komponensek)
    doc.add_heading("5. Kód dokumentáció (fő komponensek)", level=1)
    bullets = [
        "server.js: middleware-ek, route-ok, DB init, statikus fájlok.",
        "routes/*.js: auth, cars, rentals, contact, admin, pages.",
        "utils/database.js: adatbázis init, táblák, demo flotta feltöltés.",
        "public: HTML oldalak, styles.css, app.js (kliensoldali logika).",
        "AdminClient: WPF UI, AdminApiService, modellek.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    # 6. Fejlesztési lehetőségek
    doc.add_heading("6. Fejlesztési lehetőségek", level=1)
    bullets = [
        "Kétfaktoros hitelesítés (2FA).",
        "Többnyelvű támogatás (i18n).",
        "Fejlettebb riportok/grafikonok.",
        "Online fizetés integráció (Stripe/PayPal).",
        "RBAC jogosultság-kezelés, cache réteg (Redis).",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    # 7. Tesztelés
    doc.add_heading("7. Tesztelés", level=1)
    bullets = [
        "Egységtesztek: üzleti logika (foglalási ütközés, auth) Jest/Mocha.",
        "Integrációs tesztek: API endpointok (Postman/Newman vagy Jest + supertest).",
        "UI tesztek: manuális + automatizált (Cypress/Playwright) fő user flow-kra.",
        "WPF kliens: ViewModel egységteszt, manuális funkcionális teszt (login, CRUD).",
        "Teljesítmény: k6/JMeter a kritikus végpontokra.",
        "Biztonság: alap OWASP ellenőrzések, session/admin endpoint audit.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    out_path = "Fejlesztoi_dokumentacio.docx"
    doc.save(out_path)


if __name__ == "__main__":
    main()


